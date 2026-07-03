from flask import Blueprint, request, jsonify, send_file, current_app, Response
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
import json
import threading

api = Blueprint('api', __name__, url_prefix='/api')

# Global stop flag for streaming
_stop_flags = {}
_stop_lock = threading.Lock()


def register_routes(app):
    app.register_blueprint(api)


# ============== Conversations ==============

@api.route('/conversations', methods=['GET'])
def get_conversations():
    conversations = current_app.conversation_model.get_all()
    return jsonify(conversations)


@api.route('/conversations', methods=['POST'])
def create_conversation():
    data = request.json or {}
    title = data.get('title', 'New Conversation')
    conv_id = current_app.conversation_model.create(title)
    return jsonify({"id": conv_id, "title": title})


@api.route('/conversations/<id>', methods=['GET'])
def get_conversation(id):
    conv = current_app.conversation_model.get_by_id(id)
    if conv:
        return jsonify(conv)
    return jsonify({"error": "Not found"}), 404


@api.route('/conversations/<id>', methods=['DELETE'])
def delete_conversation(id):
    current_app.conversation_model.delete(id)
    return jsonify({"success": True})


@api.route('/conversations/clear', methods=['DELETE'])
def clear_all_conversations():
    current_app.conversation_model.delete_all()
    return jsonify({"success": True})


# ============== Chat ==============

@api.route('/chat', methods=['POST'])
def chat():
    from services import AIService

    data = request.json
    message = data.get('message', '')
    models = data.get('models', ['openai'])
    conversation_id = data.get('conversation_id')

    if not message:
        return jsonify({"error": "Message is required"}), 400

    # Get settings for API keys (normalize to lowercase for consistency)
    settings = current_app.settings_model.get()
    api_keys = settings.get('api_keys', {})
    # Convert keys to lowercase to match AI service expectations
    api_keys = {k.lower(): v for k, v in api_keys.items()}

    # Create AI service instance
    ai_service = AIService(api_keys)

    # Add current time as system message
    from datetime import datetime, timezone, timedelta
    utc_now = datetime.now(timezone.utc)
    cst = timezone(timedelta(hours=8))
    current_time = utc_now.astimezone(cst).strftime("%Y年%m月%d日 %H:%M:%S")
    system_message = {"role": "system", "content": f"当前日期时间(北京时间)：{current_time}"}

    results = []

    # Get existing messages for context
    existing_messages = []
    if conversation_id:
        conv = current_app.conversation_model.get_by_id(conversation_id)
        if conv:
            existing_messages = [
                {"role": m["role"], "content": m["content"]}
                for m in conv.get('messages', [])
            ]

    # Add current user message
    current_messages = [system_message] + existing_messages + [{"role": "user", "content": message}]

    # Send to each selected model
    for model in models:
        model_lower = model.lower()  # Normalize model name to lowercase
        result = ai_service.chat(model_lower, current_messages)
        # Handle both old string format and new dict format
        if isinstance(result, dict):
            content = result.get("content", "")
        else:
            content = result

        results.append({
            "model": model,  # Keep original name for display
            "content": content
        })

    # Add user message to conversation first
    if conversation_id:
        current_app.conversation_model.add_message(
            conversation_id,
            "user",
            "user",
            message
        )

    # Add assistant responses to conversation
    if conversation_id:
        for i, result in enumerate(results):
            current_app.conversation_model.add_message(
                conversation_id,
                result["model"],
                "assistant",
                result["content"]
            )

    return jsonify({
        "user_message": message,
        "responses": results
    })


@api.route('/chat/stream', methods=['POST'])
def chat_stream():
    """Streaming chat endpoint using Server-Sent Events"""
    from services import AIService
    from datetime import datetime, timezone, timedelta
    
    data = request.json
    message = data.get('message', '')
    models = data.get('models', ['openai'])
    conversation_id = data.get('conversation_id')

    if not message:
        return jsonify({"error": "Message is required"}), 400

    # Get settings and conversation data BEFORE starting the stream
    settings = current_app.settings_model.get()
    api_keys = settings.get('api_keys', {})
    api_keys = {k.lower(): v for k, v in api_keys.items()}
    
    # Capture model reference to use in generator
    conversation_model = current_app.conversation_model
    
    existing_messages = []
    if conversation_id:
        conv = conversation_model.get_by_id(conversation_id)
        if conv:
            existing_messages = [
                {"role": m["role"], "content": m["content"]}
                for m in conv.get('messages', [])
            ]
        # Add user message to conversation
        conversation_model.add_message(conversation_id, "user", "user", message)

    # Get current time
    utc_now = datetime.now(timezone.utc)
    cst = timezone(timedelta(hours=8))
    current_time = utc_now.astimezone(cst).strftime("%Y年%m月%d日 %H:%M:%S")
    system_message = {"role": "system", "content": f"当前日期时间(北京时间)：{current_time}"}

    # Add user message
    current_messages = [system_message] + existing_messages + [{"role": "user", "content": message}]

    def generate():
        ai_service = AIService(api_keys)
        
        # Process each model and stream responses
        for model in models:
            model_lower = model.lower()
            
            # Send start event for this model
            yield f"event: model_start\ndata: {json.dumps({'model': model})}\n\n"
            
            full_content = ""
            for chunk in ai_service.chat_stream(model_lower, current_messages):
                full_content += chunk
                yield f"event: content\ndata: {json.dumps({'model': model, 'content': chunk})}\n\n"
            
            # Send end event
            yield f"event: model_end\ndata: {json.dumps({'model': model, 'full_content': full_content})}\n\n"
            
            # Save to conversation
            if conversation_id:
                conversation_model.add_message(conversation_id, model, "assistant", full_content)

        yield f"event: done\ndata: {json.dumps({'message': message})}\n\n"

    def generate_with_cleanup():
        try:
            yield from generate()
        except GeneratorExit:
            # Client disconnected - cleanup will happen automatically
            pass

    return Response(
        generate_with_cleanup(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'close'
        }
    )


# ============== Documents ==============

@api.route('/documents', methods=['GET'])
def get_documents():
    docs = current_app.document_model.get_all()
    print(f"=== get_documents: returning {len(docs)} documents ===")
    print(f"Documents: {docs}")
    return jsonify(docs)


@api.route('/documents', methods=['POST'])
def create_document():
    data = request.json or {}
    title = data.get('title', 'Untitled Document')
    content = data.get('content', '')
    doc_id = current_app.document_model.create(title, content)
    return jsonify({"id": doc_id, "title": title})


@api.route('/documents/<id>', methods=['GET'])
def get_document(id):
    doc = current_app.document_model.get_by_id(id)
    print(f"=== get_document({id}) ===")
    print(f"doc: {doc}")
    if doc:
        return jsonify(doc)
    return jsonify({"error": "Not found"}), 404


@api.route('/documents/<id>', methods=['PUT'])
def update_document(id):
    data = request.json or {}
    current_app.document_model.update(
        id,
        title=data.get('title'),
        content=data.get('content')
    )
    return jsonify({"success": True})


@api.route('/documents/<id>', methods=['DELETE'])
def delete_document(id):
    current_app.document_model.delete(id)
    return jsonify({"success": True})


@api.route('/documents/<id>/export', methods=['GET'])
def export_document(id):
    doc = current_app.document_model.get_by_id(id)
    if not doc:
        return jsonify({"error": "Not found"}), 404

    # Create DOCX
    docx_doc = Document()
    docx_doc.title = doc.get('title', 'Document')

    # Parse HTML content and convert to DOCX
    html_content = doc.get('content', '')
    if html_content:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Debug: check soup structure
        if soup.body is None:
            # If no body, try to use the root element directly
            soup = BeautifulSoup(f'<body>{html_content}</body>', 'html.parser')
        
        # Helper function to check if element is a tag
        def is_tag(element):
            return hasattr(element, 'name') and element.name is not None
        
        # Process elements recursively
        def process_children(container):
            for child in container.children:
                if not is_tag(child):
                    # Text node
                    text = str(child).strip()
                    if text:
                        docx_doc.add_paragraph(text)
                    continue
                
                name = child.name
                if name in ['script', 'style']:
                    continue
                
                if name == 'p':
                    # Get text content including nested elements
                    text = child.get_text().strip()
                    if text:
                        docx_doc.add_paragraph(text)
                elif name in ['h1', 'h2', 'h3']:
                    level = int(name[1])
                    docx_doc.add_heading(child.get_text(), level=level)
                elif name == 'ul':
                    for li in child.find_all('li', recursive=False):
                        docx_doc.add_paragraph(li.get_text(), style='List Bullet')
                elif name == 'ol':
                    for li in child.find_all('li', recursive=False):
                        docx_doc.add_paragraph(li.get_text(), style='List Number')
                elif name == 'div':
                    classes = child.get('class', [])
                    if 'quill-better-table-wrapper' in classes:
                        # Find the table inside the wrapper
                        table = child.find('table', class_='quill-better-table')
                        if table:
                            rows = table.find_all('tr')
                            if rows:
                                # Count max columns
                                max_cols = 0
                                for row in rows:
                                    cells = row.find_all(['th', 'td'])
                                    max_cols = max(max_cols, len(cells))
                                
                                if max_cols > 0:
                                    # Create DOCX table
                                    docx_table = docx_doc.add_table(rows=len(rows), cols=max_cols)
                                    docx_table.style = 'Table Grid'
                                    
                                    for i, row in enumerate(rows):
                                        cells = row.find_all(['th', 'td'])
                                        for j, cell in enumerate(cells):
                                            if j < max_cols:
                                                cell_text = cell.get_text().strip()
                                                docx_table.cell(i, j).text = cell_text
                    else:
                        # Process children of other divs
                        process_children(child)
                elif name == 'table':
                    rows = child.find_all('tr')
                    if rows:
                        max_cols = 0
                        for row in rows:
                            cells = row.find_all(['th', 'td'])
                            max_cols = max(max_cols, len(cells))
                        
                        if max_cols > 0:
                            docx_table = docx_doc.add_table(rows=len(rows), cols=max_cols)
                            docx_table.style = 'Table Grid'
                            
                            for i, row in enumerate(rows):
                                cells = row.find_all(['th', 'td'])
                                for j, cell in enumerate(cells):
                                    if j < max_cols:
                                        cell_text = cell.get_text().strip()
                                        docx_table.cell(i, j).text = cell_text
                else:
                    # Process children of unknown elements
                    process_children(child)
        
        # Process body content
        if soup.body:
            process_children(soup.body)
    else:
        docx_doc.add_paragraph('')

    # Save to BytesIO
    buffer = BytesIO()
    docx_doc.save(buffer)
    buffer.seek(0)

    filename = f"{doc.get('title', 'document').replace(' ', '_')}.docx"
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


# ============== Settings ==============

@api.route('/settings', methods=['GET'])
def get_settings():
    settings = current_app.settings_model.get()
    # Return actual API keys for frontend to check
    api_keys = settings.get('api_keys', {})
    return jsonify({
        "api_keys": api_keys,
        "enabled_models": settings.get('enabled_models', []),
        "mongodb_uri": settings.get('mongodb_uri', '')
    })


@api.route('/settings', methods=['PUT'])
def update_settings():
    data = request.json or {}

    # Get existing settings to preserve API keys if not provided
    existing = current_app.settings_model.get()
    existing_keys = existing.get('api_keys', {})

    # Merge API keys - only update if value is a non-empty string
    new_keys = data.get('api_keys', {})
    for k, v in new_keys.items():
        if isinstance(v, str) and v.strip():  # Only update if value is a valid string
            existing_keys[k] = v

    # Get enabled models
    enabled_models = data.get('enabled_models', [])

    current_app.settings_model.update(
        api_keys=existing_keys,
        enabled_models=enabled_models,
        mongodb_uri=data.get('mongodb_uri')
    )
    return jsonify({"success": True})


@api.route('/models', methods=['GET'])
def get_models():
    """Return list of available AI models"""
    return jsonify([
        {"id": "OpenAI", "name": "ChatGPT", "icon": "/static/icons/chatgpt.svg"},
        {"id": "Claude", "name": "Claude", "icon": "/static/icons/claude.svg"},
        {"id": "Doubao", "name": "Doubao-Seed", "icon": "/static/icons/doubao.svg"},
        {"id": "DeepSeek", "name": "DeepSeek", "icon": "/static/icons/deepseek.svg"},
        {"id": "Gemini", "name": "Gemini", "icon": "/static/icons/gemini.svg"},
        {"id": "GLM", "name": "GLM", "icon": "/static/icons/zhipu.svg"},
        {"id": "Kimi", "name": "Kimi", "icon": "/static/icons/kimi.svg"},
        {"id": "MiniMax", "name": "MiniMax", "icon": "/static/icons/MiniMax.svg"},
        {"id": "Qwen", "name": "Qwen", "icon": "/static/icons/Qianwen.svg"},
        {"id": "Yuanbao", "name": "Yuanbao", "icon": "/static/icons/yuanbao.svg"}
    ])
